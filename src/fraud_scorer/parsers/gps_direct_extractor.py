"""Advanced GPS direct extractor with normalization helpers."""

from __future__ import annotations

import hashlib
import logging
import math
import os
import re
import time
import unicodedata
import importlib
from dataclasses import dataclass, field
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Dict, Iterable, List, Optional, Sequence, Tuple

import pandas as pd

from fraud_scorer.parsers.processing_hint import ProcessingHint
from fraud_scorer.parsers.types import ParsedDocument, Table


logger = logging.getLogger(__name__)


GPS_SCHEMA_VERSION = 1
REQUIRED_COLUMNS = (
    "timestamp",
    "latitude",
    "longitude",
    "speed",
    "heading",
    "event_label",
)
OPTIONAL_COLUMNS = (
    "odometer",
    "fuel_level",
    "driver_id",
    "source_page",
    "source_plugin",
)
RAW_PAYLOAD_COLUMN = "raw_row_json"

# Column aliases normalizados (minúsculas sin acentos ni espacios)
COLUMN_ALIASES: Dict[str, set[str]] = {
    "timestamp": {
        "timestamp",
        "time",
        "datetime",
        "fecha",
        "fechahora",
        "fecha_hora",
        "fechaevento",
        "fechayhora",
        "hora",
        "fechayhoraevento",
        "eventtime",
    },
    "latitude": {"latitude", "lat", "latitud", "coordlat", "coordenadalat"},
    "longitude": {"longitude", "lon", "long", "longitud", "coordlon", "coordenadalon"},
    "speed": {"speed", "velocidad", "vel", "velocidadkmh", "kmh", "mph"},
    "heading": {"heading", "rumbo", "direccion", "bearing"},
    "event_label": {"event", "evento", "eventlabel", "estado", "status", "description"},
    "odometer": {"odometer", "odometro", "odometrah", "kms", "kilometraje"},
    "fuel_level": {"fuellevel", "combustible", "nivelcombustible", "fuel"},
    "driver_id": {"driver", "conductor", "operador", "driverid", "idoperador"},
}


@dataclass
class PluginResult:
    tables: List[Table] = field(default_factory=list)
    text_segments: List[str] = field(default_factory=list)
    page_stats: List[Dict[str, Any]] = field(default_factory=list)
    plugin_name: str = "unknown"
    extra_metadata: Dict[str, Any] = field(default_factory=dict)
    warnings: List[str] = field(default_factory=list)


class GPSExtractionPlugin:
    """Interface for extractor plugins."""

    extensions: Sequence[str] = ()
    name: str = "base"

    def supports(self, extension: str) -> bool:
        return extension.lower() in {ext.lower() for ext in self.extensions}

    def extract(self, path: Path, hint: Optional[ProcessingHint]) -> PluginResult:  # pragma: no cover - virtual
        raise NotImplementedError


class PDFExtractionPlugin(GPSExtractionPlugin):
    extensions = (".pdf",)
    name = "pdf-direct"

    def extract(self, path: Path, hint: Optional[ProcessingHint]) -> PluginResult:
        import fitz  # PyMuPDF

        text_segments: List[str] = []
        page_stats: List[Dict[str, Any]] = []
        tables: List[Table] = []
        warnings: List[str] = []
        plugins_used: List[str] = []

        doc = fitz.open(path)
        try:
            for page_index, page in enumerate(doc, start=1):
                text = page.get_text("text") or ""
                text_segments.append(text)
                images = page.get_images(full=True)
                page_stats.append(
                    {
                        "page": page_index,
                        "char_count": len(text),
                        "image_count": len(images),
                    }
                )
        finally:
            doc.close()

        # Primer intento con pdfplumber
        try:
            pdfplumber = importlib.import_module("pdfplumber")
        except ModuleNotFoundError:
            warnings.append("pdfplumber_no_disponible")
        except Exception as exc:  # pragma: no cover - fallback
            warnings.append(f"pdfplumber_import_error:{exc}")
        else:
            try:
                with pdfplumber.open(path) as pdf:
                    for page_index, page in enumerate(pdf.pages, start=1):
                        extracted_tables = page.extract_tables() or []
                        for table_index, raw_table in enumerate(extracted_tables, start=1):
                            if not raw_table:
                                continue
                            headers = [clean_cell(cell) for cell in (raw_table[0] or [])]
                            rows = [
                                [clean_cell(cell) for cell in (row or [])]
                                for row in raw_table[1:]
                                if any(clean_cell(cell) for cell in (row or []))
                            ]
                            if not rows:
                                continue
                            tables.append(
                                {
                                    "headers": headers,
                                    "data_rows": rows,
                                    "sheet_name": f"page_{page_index}",
                                    "row_count": len(rows),
                                    "column_count": len(headers),
                                    "source_page": page_index,
                                    "gps_plugin": "pdfplumber",
                                }
                            )
                if tables:
                    plugins_used.append("pdfplumber")
            except Exception as exc:  # pragma: no cover - fallback
                warnings.append(f"pdfplumber_error:{exc}")

        # Fallback con Camelot si no se extrajeron tablas suficientes
        if not tables:
            try:
                camelot = importlib.import_module("camelot")
            except ModuleNotFoundError:
                warnings.append("camelot_no_disponible")
            except Exception as exc:  # pragma: no cover - fallback
                warnings.append(f"camelot_import_error:{exc}")
            else:
                try:
                    camelot_tables = camelot.read_pdf(str(path), pages="all", flavor="stream")
                    for tbl in camelot_tables:
                        df = tbl.df
                        headers = [clean_cell(col) or f"col_{idx+1}" for idx, col in enumerate(df.iloc[0].tolist())]
                        data_rows = [
                            [clean_cell(value) for value in row]
                            for row in df.iloc[1:].values.tolist()
                        ]
                        if not data_rows:
                            continue
                        tables.append(
                            {
                                "headers": headers,
                                "data_rows": data_rows,
                                "sheet_name": f"camelot_{tbl.page}",
                                "row_count": len(data_rows),
                                "column_count": len(headers),
                                "source_page": getattr(tbl, "page", None),
                                "gps_plugin": "camelot",
                            }
                        )
                    if camelot_tables:
                        plugins_used.append("camelot")
                except Exception as exc:  # pragma: no cover - fallback
                    warnings.append(f"camelot_error:{exc}")

        return PluginResult(
            tables=tables,
            text_segments=text_segments,
            page_stats=page_stats,
            plugin_name=self.name,
            extra_metadata={"plugins_used": plugins_used},
            warnings=warnings,
        )


class CSVExtractionPlugin(GPSExtractionPlugin):
    extensions = (".csv",)
    name = "csv"

    def __init__(self, *, chunk_rows: Optional[int] = None, preview_rows: int = 500) -> None:
        self.chunk_rows = chunk_rows
        self.preview_rows = preview_rows

    def extract(self, path: Path, hint: Optional[ProcessingHint]) -> PluginResult:
        chunk_rows = self.chunk_rows
        reader = pd.read_csv(
            path,
            sep=None,
            engine="python",
            encoding="utf-8",
            dtype=str,
            chunksize=chunk_rows,
        )

        tables: List[Table] = []
        text_segments: List[str] = []

        if isinstance(reader, pd.DataFrame):
            df = reader.where(pd.notnull(reader), None)
            tables.append(dataframe_to_table(df, sheet_name="csv", gps_plugin=self.name))
            text_segments.append(df.head(self.preview_rows).to_string(index=False))
        else:
            for index, chunk in enumerate(reader, start=1):
                cleaned = chunk.where(pd.notnull(chunk), None)
                tables.append(
                    dataframe_to_table(
                        cleaned,
                        sheet_name=f"csv_chunk_{index}",
                        gps_plugin=self.name,
                    )
                )
                if index == 1:
                    text_segments.append(cleaned.head(self.preview_rows).to_string(index=False))

        return PluginResult(
            tables=tables,
            text_segments=text_segments or [""],
            plugin_name=self.name,
            extra_metadata={"chunk_count": len(tables)},
        )


class ExcelExtractionPlugin(GPSExtractionPlugin):
    extensions = (".xlsx", ".xls")
    name = "excel"

    def __init__(self, *, chunk_rows: Optional[int] = None, preview_rows: int = 200) -> None:
        self.chunk_rows = chunk_rows
        self.preview_rows = preview_rows

    def extract(self, path: Path, hint: Optional[ProcessingHint]) -> PluginResult:
        xls = pd.ExcelFile(path)
        tables: List[Table] = []
        text_segments: List[str] = []
        for sheet in xls.sheet_names:
            df = pd.read_excel(xls, sheet_name=sheet, dtype=str)
            df = df.where(pd.notnull(df), None)

            if self.chunk_rows and df.shape[0] > self.chunk_rows:
                for idx, chunk in enumerate(_yield_dataframe_chunks(df, self.chunk_rows), start=1):
                    tables.append(
                        dataframe_to_table(
                            chunk,
                            sheet_name=f"{sheet}_part_{idx}",
                            gps_plugin=self.name,
                        )
                    )
                    if idx == 1:
                        text_segments.append(
                            f"--- Hoja: {sheet} (segmento 1) ---\n{chunk.head(self.preview_rows).to_string(index=False)}"
                        )
            else:
                tables.append(dataframe_to_table(df, sheet_name=sheet, gps_plugin=self.name))
                text_segments.append(f"--- Hoja: {sheet} ---\n{df.head(self.preview_rows).to_string(index=False)}")
        return PluginResult(
            tables=tables,
            text_segments=text_segments,
            plugin_name=self.name,
            extra_metadata={"chunk_count": len(tables)},
        )


class DocxExtractionPlugin(GPSExtractionPlugin):
    extensions = (".docx",)
    name = "docx"

    def extract(self, path: Path, hint: Optional[ProcessingHint]) -> PluginResult:
        import docx  # type: ignore

        document = docx.Document(path)
        text_parts = [p.text for p in document.paragraphs if p.text]
        tables: List[Table] = []
        for t in document.tables:
            headers = [cell.text for cell in t.rows[0].cells] if t.rows else []
            rows: List[List[Any]] = []
            for row in t.rows[1:]:
                rows.append([cell.text for cell in row.cells])
            tables.append(
                {
                    "headers": headers,
                    "data_rows": rows,
                    "sheet_name": "docx",
                    "row_count": len(rows),
                    "column_count": len(headers),
                    "source_page": None,
                    "gps_plugin": self.name,
                }
            )
        return PluginResult(
            tables=tables,
            text_segments=text_parts,
            plugin_name=self.name,
            extra_metadata={"chunk_count": len(tables)},
        )


class GPSDirectExtractor:
    """Transforms native GPS sources into structured outputs and metadata."""

    SUPPORTED_EXT = {".pdf", ".csv", ".xlsx", ".xls", ".docx"}

    def __init__(
        self,
        *,
        max_file_mb: float = 500.0,
        csv_chunk_rows: Optional[int] = None,
        excel_chunk_rows: Optional[int] = None,
    ) -> None:
        self.max_file_bytes = int(max_file_mb * 1024 * 1024)
        csv_chunks = csv_chunk_rows or _default_chunk_rows()
        excel_chunks = excel_chunk_rows or _default_chunk_rows()
        self.plugins: Sequence[GPSExtractionPlugin] = (
            PDFExtractionPlugin(),
            CSVExtractionPlugin(chunk_rows=csv_chunks),
            ExcelExtractionPlugin(chunk_rows=excel_chunks),
            DocxExtractionPlugin(),
        )

    def extract(self, path: Path, hint: Optional[ProcessingHint] = None) -> ParsedDocument:
        extension = path.suffix.lower()
        if extension not in self.SUPPORTED_EXT:
            raise ValueError(f"Formato no soportado por GPSDirectExtractor: {extension}")
        if hint and hint.file_size_bytes and hint.file_size_bytes > self.max_file_bytes:
            raise ValueError(
                f"Archivo excede el máximo permitido para ingesta directa ({hint.file_size_bytes} > {self.max_file_bytes})"
            )

        plugin = self._resolve_plugin(extension)
        start = time.perf_counter()
        result = plugin.extract(path, hint)
        processing_ms = int((time.perf_counter() - start) * 1000)

        metadata = self._base_metadata(path, plugin.name, hint)

        df, normalization_warnings = normalize_gps_tables(result.tables)
        preview_rows = _serialize_preview_rows(df.head(20)) if not df.empty else []
        summary = build_gps_summary(
            df,
            page_stats=result.page_stats,
            extra_warnings=result.warnings + normalization_warnings,
        )

        plugins_used = [plugin.name]
        plugins_used.extend(result.extra_metadata.get("plugins_used", []))
        chunk_count = result.extra_metadata.get("chunk_count")

        metadata.setdefault("gps_direct", {})
        metadata["gps_direct"].update(
            {
                "enabled": True,
                "hint": hint.as_dict() if hint else None,
                "schema_version": GPS_SCHEMA_VERSION,
                "normalized_row_count": int(df.shape[0]),
                "columns_detected": list(df.columns),
                "plugins_used": plugins_used,
                "normalization_warnings": normalization_warnings,
                "ingestion_stats": {
                    "processing_time_ms": processing_ms,
                    "table_count": len(result.tables),
                    "page_count": len(result.page_stats),
                    "dataframe_columns": list(df.columns),
                    "dataframe_rows": int(df.shape[0]),
                },
                "preview_rows": preview_rows,
            }
        )
        if chunk_count is not None:
            metadata["gps_direct"]["chunk_count"] = int(chunk_count)

        gps_summary = {
            **summary,
            "generated_at": datetime.now(timezone.utc).isoformat(),
        }

        return {
            "text": "\n".join(result.text_segments),
            "tables": result.tables,
            "key_value_pairs": {},
            "metadata": metadata,
            "gps_summary": gps_summary,
        }

    # ------------------------------------------------------------------
    # Helpers
    # ------------------------------------------------------------------
    def _resolve_plugin(self, extension: str) -> GPSExtractionPlugin:
        for plugin in self.plugins:
            if plugin.supports(extension):
                return plugin
        raise ValueError(f"No existe plugin para la extensión {extension}")

    def _base_metadata(
        self,
        path: Path,
        source_type: str,
        hint: Optional[ProcessingHint],
    ) -> Dict[str, Any]:
        try:
            file_size = path.stat().st_size
        except Exception:
            file_size = 0
        metadata: Dict[str, Any] = {
            "source_type": f"gps_direct_{source_type}",
            "file_name": path.name,
            "file_size": file_size,
            "engine": "gps_dal",
        }
        if hint:
            metadata.setdefault("gps_direct", {})
            metadata["gps_direct"]["hint"] = hint.as_dict()
        return metadata


# ----------------------------------------------------------------------
# Normalization helpers (reused by GPSCacheManager)
# ----------------------------------------------------------------------


def normalize_gps_tables(tables: Iterable[Table]) -> Tuple[pd.DataFrame, List[str]]:
    """Normalizes raw tables into the canonical GPS schema."""

    records: List[Dict[str, Any]] = []
    warnings: List[str] = []

    for table in tables:
        headers = list(table.get("headers") or [])
        rows = list(table.get("data_rows") or [])
        if not rows:
            continue
        normalized_headers = _normalize_headers(headers, len(rows[0]))
        page = table.get("source_page")
        plugin = table.get("gps_plugin")

        for row in rows:
            padded_row = _pad_row(row, len(normalized_headers))
            raw_row = {
                normalized_headers[idx]: padded_row[idx]
                for idx in range(len(normalized_headers))
            }
            canonical_row = _map_row_to_canonical(raw_row)
            canonical_row.setdefault("source_page", page)
            canonical_row.setdefault("source_plugin", plugin)
            canonical_row[RAW_PAYLOAD_COLUMN] = raw_row
            records.append(canonical_row)

    if not records:
        return pd.DataFrame(columns=REQUIRED_COLUMNS + OPTIONAL_COLUMNS + (RAW_PAYLOAD_COLUMN,)), warnings

    df = pd.DataFrame(records)

    # Ensure all expected columns exist
    for column in REQUIRED_COLUMNS + OPTIONAL_COLUMNS:
        if column not in df.columns:
            df[column] = None

    # Type coercions
    df["timestamp"] = df["timestamp"].apply(_coerce_timestamp)
    for column in ("latitude", "longitude", "speed", "heading", "odometer", "fuel_level"):
        if column in df.columns:
            df[column] = df[column].apply(_coerce_float)

    # Normalize event label string
    if "event_label" in df.columns:
        df["event_label"] = df["event_label"].apply(
            lambda value: clean_cell(value) or None
        )

    df["driver_id"] = df["driver_id"].apply(lambda value: clean_cell(value) or None)
    df["source_plugin"] = df["source_plugin"].apply(lambda value: clean_cell(value) or None)

    df.sort_values(by=["timestamp", "latitude", "longitude"], inplace=True, na_position="last")
    df.reset_index(drop=True, inplace=True)

    missing_coords = df[(df["latitude"].isna()) | (df["longitude"].isna())]
    if not missing_coords.empty:
        warnings.append("missing_coordinates")

    if df["timestamp"].isna().all():
        warnings.append("missing_timestamps")

    # Serialize raw rows as JSON strings to keep dataset lightweight
    df[RAW_PAYLOAD_COLUMN] = df[RAW_PAYLOAD_COLUMN].apply(_serialize_raw_row)

    ordered_columns = list(REQUIRED_COLUMNS) + list(OPTIONAL_COLUMNS) + [RAW_PAYLOAD_COLUMN]
    df = df[ordered_columns]
    return df, warnings


def build_gps_summary(
    df: pd.DataFrame,
    *,
    page_stats: Optional[List[Dict[str, Any]]] = None,
    extra_warnings: Optional[Sequence[str]] = None,
) -> Dict[str, Any]:
    summary: Dict[str, Any] = {
        "row_count": int(df.shape[0]),
        "columns": list(df.columns),
        "warnings": list(extra_warnings or []),
    }
    if page_stats:
        summary["page_stats"] = page_stats

    if df.empty:
        return summary

    timestamps = df["timestamp"].dropna()
    if not timestamps.empty:
        summary["time_span"] = {
            "start": timestamps.min().isoformat(),
            "end": timestamps.max().isoformat(),
            "duration_minutes": int((timestamps.max() - timestamps.min()).total_seconds() / 60),
        }
        gaps = _detect_time_gaps(timestamps.sort_values())
        if gaps:
            summary["time_gaps"] = gaps

    latitudes = df["latitude"].dropna()
    longitudes = df["longitude"].dropna()
    if not latitudes.empty and not longitudes.empty:
        summary["bounding_box"] = {
            "min_lat": float(latitudes.min()),
            "max_lat": float(latitudes.max()),
            "min_lon": float(longitudes.min()),
            "max_lon": float(longitudes.max()),
        }
        try:
            h3 = importlib.import_module("h3")
        except ModuleNotFoundError:
            pass
        except Exception:  # pragma: no cover - opcional
            pass
        else:
            try:
                cells = {
                    h3.geo_to_h3(lat, lon, 7)
                    for lat, lon in zip(latitudes, longitudes)
                    if not (math.isnan(lat) or math.isnan(lon))
                }
                summary["h3_cells"] = list(cells)
            except Exception:  # pragma: no cover - opcional
                pass

    if "speed" in df.columns:
        speed_series = df["speed"].dropna()
        if not speed_series.empty:
            summary["speed_stats"] = {
                "avg": float(speed_series.mean()),
                "p95": float(speed_series.quantile(0.95)),
                "max": float(speed_series.max()),
            }

    return summary


def dataframe_to_table(df: pd.DataFrame, *, sheet_name: str, gps_plugin: str) -> Table:
    df = df.where(pd.notnull(df), None)
    return {
        "headers": list(df.columns),
        "data_rows": df.values.tolist(),
        "sheet_name": sheet_name,
        "row_count": int(df.shape[0]),
        "column_count": int(df.shape[1]),
        "source_page": sheet_name,
        "gps_plugin": gps_plugin,
}


def _yield_dataframe_chunks(df: pd.DataFrame, chunk_rows: int) -> Iterable[pd.DataFrame]:
    if chunk_rows <= 0:
        yield df
        return
    total_rows = df.shape[0]
    for start in range(0, total_rows, chunk_rows):
        yield df.iloc[start : start + chunk_rows].copy()


def _default_chunk_rows() -> int:
    try:
        return int(os.getenv("GPS_DIRECT_CHUNK_ROWS", "200000"))
    except Exception:
        return 200000


def clean_cell(value: Any) -> str:
    if value is None:
        return ""
    return str(value).strip()


def _normalize_headers(headers: List[str], first_row_len: int) -> List[str]:
    normalized = [clean_cell(header) for header in headers]
    header_count = len(normalized)
    if header_count < first_row_len:
        for idx in range(header_count, first_row_len):
            normalized.append(f"column_{idx + 1}")
    elif header_count > first_row_len:
        normalized = normalized[:first_row_len]
    return [header or f"column_{idx + 1}" for idx, header in enumerate(normalized)]


def _pad_row(row: Sequence[Any], length: int) -> List[Any]:
    padded = list(row)
    if len(padded) < length:
        padded.extend([None] * (length - len(padded)))
    elif len(padded) > length:
        padded = padded[:length]
    return padded


def _map_row_to_canonical(raw_row: Dict[str, Any]) -> Dict[str, Any]:
    canonical: Dict[str, Any] = {}
    for column, value in raw_row.items():
        normalized = _normalize_key(column)
        target = None
        for canonical_name, aliases in COLUMN_ALIASES.items():
            if normalized in aliases:
                target = canonical_name
                break
        if target:
            canonical[target] = value
    return canonical


def _normalize_key(value: str) -> str:
    normalized = unicodedata.normalize("NFKD", value)
    normalized = "".join(char for char in normalized if not unicodedata.combining(char))
    normalized = normalized.lower()
    normalized = re.sub(r"[^a-z0-9]", "", normalized)
    return normalized


def _coerce_timestamp(value: Any) -> Optional[datetime]:
    if value in (None, "", "nan", "NA"):
        return None
    try:
        ts = pd.to_datetime(value, errors="coerce")
        if pd.isna(ts):
            return None
        if ts.tzinfo is None:
            ts = ts.tz_localize(timezone.utc)
        else:
            ts = ts.astimezone(timezone.utc)
        return ts.to_pydatetime()
    except Exception:
        return None


_FLOAT_PATTERN = re.compile(r"[^0-9,.-]+")


def _coerce_float(value: Any) -> Optional[float]:
    if value in (None, "", "nan", "NA"):
        return None
    value_str = str(value).strip()
    value_str = _FLOAT_PATTERN.sub("", value_str)
    if value_str.count(",") == 1 and value_str.count(".") == 0:
        value_str = value_str.replace(",", ".")
    else:
        value_str = value_str.replace(",", "")
    try:
        return float(value_str)
    except Exception:
        return None


def _serialize_raw_row(raw_row: Dict[str, Any]) -> str:
    try:
        import orjson

        return orjson.dumps(raw_row, option=orjson.OPT_NON_STR_KEYS).decode("utf-8")
    except Exception:
        import json

        return json.dumps(raw_row, ensure_ascii=False, default=str)


def _serialize_preview_rows(df: pd.DataFrame) -> List[Dict[str, Any]]:
    safe_df = df.copy()
    if "timestamp" in safe_df.columns:
        safe_df["timestamp"] = safe_df["timestamp"].apply(
            lambda ts: ts.isoformat()
            if isinstance(ts, datetime)
            else (ts.isoformat() if hasattr(ts, "isoformat") else ts)
        )
    safe_df = safe_df.where(pd.notnull(safe_df), None)
    return safe_df.to_dict(orient="records")


def _detect_time_gaps(timestamps: pd.Series) -> List[Dict[str, Any]]:
    gaps: List[Dict[str, Any]] = []
    diffs = timestamps.diff().dropna()
    threshold = pd.Timedelta(hours=1)
    indices = diffs[diffs > threshold].index
    for idx in indices:
        previous_ts = timestamps.loc[idx - 1]
        current_ts = timestamps.loc[idx]
        gaps.append(
            {
                "start": previous_ts.isoformat(),
                "end": current_ts.isoformat(),
                "gap_minutes": int((current_ts - previous_ts).total_seconds() / 60),
            }
        )
    return gaps


def compute_sha256_for_paths(paths: Iterable[Path]) -> str:
    hasher = hashlib.sha256()
    for path in sorted(paths):
        try:
            hasher.update(path.name.encode("utf-8"))
            with open(path, "rb") as file_handler:
                while True:
                    chunk = file_handler.read(1024 * 1024)
                    if not chunk:
                        break
                    hasher.update(chunk)
        except FileNotFoundError:
            continue
    return hasher.hexdigest()


__all__ = [
    "GPSDirectExtractor",
    "normalize_gps_tables",
    "build_gps_summary",
    "dataframe_to_table",
    "compute_sha256_for_paths",
    "GPS_SCHEMA_VERSION",
    "RAW_PAYLOAD_COLUMN",
]
