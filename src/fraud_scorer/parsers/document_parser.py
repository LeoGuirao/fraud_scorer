# src/fraud_scorer/parsers/document_parser.py
"""
DocumentParser: Orquestador para procesar múltiples formatos de documentos.

Este módulo identifica el tipo de archivo y delega el procesamiento al
lector/procesador correspondiente, devolviendo SIEMPRE una salida unificada:
{
  "text": str,
  "tables": List[Table],
  "key_value_pairs": Dict[str, Any],
  "metadata": DocumentMetadata
}
"""

from __future__ import annotations

import logging
from pathlib import Path
from typing import Any, Dict, Optional

# Lectores y tipos unificados
from fraud_scorer.parsers.types import (
    DocumentReader,
    ParsedDocument,
    OCR_EXTENSIONS,
)
from fraud_scorer.parsers.readers.azure_reader import AzureOCRReader
from fraud_scorer.parsers.document_router import DocumentIntakeRouter
from fraud_scorer.parsers.processing_hint import (
    ProcessingHint,
    ProcessingHintBuilder,
)

# OCR de Azure (inyectado al adapter)
from fraud_scorer.processors.ocr.azure_ocr import AzureOCRProcessor

# Librerías para formatos nativos
import docx
import pandas as pd


logger = logging.getLogger(__name__)


class DocumentParser:
    """
    Orquesta el parsing de diferentes tipos de documentos.

    - Para imágenes/PDF: usa un lector OCR (adapter) que normaliza la salida.
    - Para DOCX/XLSX/CSV: usa parsers nativos y normaliza aquí mismo.
    """

    def __init__(self, ocr_processor: AzureOCRProcessor):
        """
        Inicializa el parser con los procesadores necesarios.

        Args:
            ocr_processor: Instancia del procesador OCR (AzureOCRProcessor).
        """
        # Adapter/Reader que encapsula Azure y devuelve salida unificada
        self.ocr_reader: DocumentReader = AzureOCRReader(ocr_processor)
        self.hint_builder = ProcessingHintBuilder()
        self.intake_router = DocumentIntakeRouter(hint_builder=self.hint_builder)
        logger.info("DocumentParser inicializado con todos los procesadores.")

    def parse_document(
        self,
        doc_path: Path,
        *,
        hint: Optional[ProcessingHint] = None,
    ) -> Optional[ParsedDocument]:
        """
        Parsea un documento, seleccionando el método apropiado según su extensión.

        Args:
            doc_path: Ruta al documento a procesar.

        Returns:
            ParsedDocument con salida unificada o None si no es soportado o hay error.
        """
        if not doc_path.exists():
            logger.error(f"El archivo no existe: {doc_path}")
            return None

        # Ignorar archivos ocultos temporales (p. ej., ._archivo.pdf en macOS)
        if doc_path.name.startswith("._"):
            logger.warning(f"Omitiendo archivo temporal/oculto: {doc_path.name}")
            return None

        logger.info("Iniciando parsing para: %s", doc_path.name)

        router_hint = hint or self.hint_builder.build(doc_path)
        try:
            result = self.intake_router.route(doc_path, self._legacy_parse, hint=router_hint)
            if result and router_hint:
                metadata = result.setdefault("metadata", {})
                metadata["processing_hint"] = router_hint.as_dict()
                gps_meta = metadata.setdefault("gps_direct", {}) if metadata.get("gps_direct") else None
                if gps_meta is not None and not gps_meta.get("hint"):
                    gps_meta["hint"] = router_hint.as_dict()
            return result
        except Exception as exc:
            logger.error("Error en intaked routing para %s: %s", doc_path.name, exc, exc_info=True)
            return None

    # ==========================
    # Parsers nativos unificados
    # ==========================

    def _legacy_parse(self, doc_path: Path) -> Optional[ParsedDocument]:
        """Mantiene el flujo tradicional basado en OCR/parsers nativos."""
        ext = doc_path.suffix.lower()
        try:
            if ext in OCR_EXTENSIONS:
                return self.ocr_reader.read(doc_path)

            if ext == ".docx":
                return self._parse_docx(doc_path)

            if ext == ".xlsx":
                return self._parse_excel(doc_path)

            if ext == ".csv":
                return self._parse_csv(doc_path)

            logger.warning(f"Formato no soportado: {ext} → {doc_path.name}")
            return None
        except Exception as exc:
            logger.error(f"Error al parsear {doc_path.name}: {exc}", exc_info=True)
            return None

    def _parse_docx(self, doc_path: Path) -> ParsedDocument:
        """Parsea un archivo .docx a la salida unificada."""
        document = docx.Document(doc_path)
        full_text = "\n".join(p.text for p in document.paragraphs if p.text is not None)

        tables = []
        for t in document.tables:
            # Headers seguros (si hay filas)
            headers = [c.text for c in t.rows[0].cells] if t.rows and len(t.rows) > 0 else []
            data_rows = []
            if t.rows and len(t.rows) > 1:
                for r in t.rows[1:]:
                    data_rows.append([c.text for c in r.cells])

            tables.append({
                "headers": headers,
                "data_rows": data_rows,
            })

        return {
            "text": full_text,
            "tables": tables,
            "key_value_pairs": {},  # K-V se puede inferir luego con IA
            "metadata": {
                "source_type": "docx",
                "file_name": doc_path.name,
            },
        }

    def _parse_excel(self, doc_path: Path) -> ParsedDocument:
        """Parses .xlsx a salida unificada (cada hoja → tabla)."""
        xls = pd.ExcelFile(doc_path)
        text_parts = []
        tables = []

        for sheet in xls.sheet_names:
            df = pd.read_excel(xls, sheet_name=sheet)
            # Texto completo de la hoja
            text_parts.append(f"--- Hoja: {sheet} ---\n{df.to_string(index=False)}")
            # Tabla normalizada
            tables.append({
                "sheet_name": sheet,
                "headers": df.columns.tolist(),
                "data_rows": df.where(pd.notnull(df), None).values.tolist(),  # NaN → None
            })

        return {
            "text": "\n\n".join(text_parts),
            "tables": tables,
            "key_value_pairs": {},
            "metadata": {
                "source_type": "xlsx",
                "file_name": doc_path.name,
                "sheets": xls.sheet_names,
            },
        }

    def _parse_csv(self, doc_path: Path) -> ParsedDocument:
        """Parses .csv a salida unificada (texto + tabla)."""
        try:
            df = pd.read_csv(doc_path, sep=None, engine="python", encoding="utf-8")
        except (UnicodeDecodeError, pd.errors.ParserError):
            logger.warning(f"No se pudo decodificar {doc_path.name} con UTF-8; intentando con latin1.")
            df = pd.read_csv(doc_path, sep=None, engine="python", encoding="latin1")

        return {
            "text": df.to_string(index=False),
            "tables": [{
                "headers": df.columns.tolist(),
                "data_rows": df.where(pd.notnull(df), None).values.tolist(),
            }],
            "key_value_pairs": {},
            "metadata": {
                "source_type": "csv",
                "file_name": doc_path.name,
            },
        }
